import Reader
from docx import Document, ImagePart
from os.path import basename #返回文件名 https://www.runoob.com/python/python-os-path.html
import os, shutil
from pathlib import Path
from datetime import datetime



class WordReader:

    def __init__(self, src: str, office_list: [], imgdir: str) -> None:
        self.src = src
        self.office_list = office_list
        self.ImageDir = imgdir
        
    
    def getText(self):  #文本
        res_text = []
        c = 0
        res_img  = []
        img_name = []
        res_table = []
        for file in self.office_list:
            if file.split('.')[-1] == 'docx' :
                document = Document(file)
                tables = document.tables #表格集
                for paragraph in document.paragraphs:
                    res_text.append(paragraph.text+"\n")
                    imgs = paragraph._element.xpath('.//pic:pic') #获取所有图片
                    for img in imgs: #https://blog.csdn.net/qq_39147299/article/details/125544621
                        for img_id in img.xpath('.//a:blip/@r:embed'): #获取图片id
                            part = document.part.related_parts[img_id] #根据图片id获取对应的图片
                            if isinstance(part, ImagePart): #保存图片
                                res_img.append(part.blob)
                                file = file.split('/')[-1]
                                # img_name.append('../IMAGE/'+file+'_'+basename(part.partname))
                                suffix = str(c)+'.png'
                                c += 1
                                img_name.append(self.ImageDir +file+'_'+suffix)
                tables = document.tables #表格集
                for table in document.tables:
                    cells = table._cells
                    cells_string = [cell.text for cell in cells]
                    res_table.append(cells_string)
        with open("fileDIR/word_text.txt", 'w') as f:
            for i in res_text:
                f.write(i)
            for i in res_table:
                f.write(i)
        Path('../IMAGE').mkdir(parents=True,exist_ok=True)
        c = 0
        for i in img_name: #保存图片
            # i = i.split('.')[0] + '.png'
            with open(i, "wb") as f:
                f.write(res_img[c])
            c += 1

if __name__ == '__main__':
    WR = Reader.Reader('dDIR/')
    WR.file_reader()
    print(WR.office_list)
    wr = WordReader("dDIR/", WR.office_list)
    wr.getText()
    # wr.getPicture()