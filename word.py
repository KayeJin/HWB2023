import Reader
from docx import Document, ImagePart
from os.path import basename #返回文件名 https://www.runoob.com/python/python-os-path.html

class WordReader:

    def __init__(self, src: str, office_list: []) -> None:
        self.src = src
        self.office_list = office_list
        

    def getText(self):  #文本
        res = []
        print(self.office_list)
        print(self.src)
        for file in self.office_list:
            if file.split('.')[1] == 'docx' :
                document = Document(self.src+'/'+file)
                for paragraph in document.paragraphs:
                    res.append(paragraph.text+"\n")
                    # print(paragraph.text)
                # document.close()
        with open("../word_text", 'w') as f:
            for i in res:
                f.write(i)
    
    def getTable(self): # https://blog.csdn.net/zhouz92/article/details/107179616
        for file in self.office_list:
            if file.split('.')[1] == 'docx':
                document = Document(self.src+'/'+file)
                tables = document.tables #表格集
                for table in tables:
                    cells = table._cells
                    cells_string = [cell.text for cell in cells]
                    print(cells_string) #待写入文件

    def getPicture(self, document, paragraph) -> []:
        res = []
        imgs = paragraph._element.xpath('.//pic:pic') #获取所有图片
        if not img:
            return []
        # img = img[0]
        # embed = img.xpath('.//a:blip/@r:embed')[0]
        # related_part = document.part.related_parts[embed]
        # image = related_part.image
        # res.append(image)
        for img in imgs: #https://blog.csdn.net/qq_39147299/article/details/125544621
            for img_id in img.xpath('.//a:blip/@r:embed'): #获取图片id
                part = document.part.related_parts[img_id] #根据图片id获取对应的图片
                if isinstance(part, ImagePart): #保存图片
                    with open(basename(part.partname), "wb") as f:
                        f.write(part.blob)


    # def showPicture(self):
    #     for file in self.office_list:
    #         if file.split('.')[1] == 'docx':
    #             document = Document(self.src+'/'+file)
    #             for paragraph in document.paragraphs:
    #                 img_list = self.getPicture(document, paragraph)
    #     with open("../image_text", 'w') as f:
    #         for img in img_list:
    #             blob = img.blob
    #             f.write(blob)
if __name__ == '__main__':
    WR = Reader.Reader('../赛题材料/wps')
    WR.file_reader()
    wr = WordReader("../赛题材料/wps", WR.office_list)
    wr.getText()
